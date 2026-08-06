import os

from wpipe import step, to_obj

from ..dto.post_train_context import PostTrainContext
from ..utils.training_report_analyzer import TrainingReportAnalyzer


@step(name="llm_analyzer", version="v1.0")
class LlmAnalyzer:

    RESULTS_RELATIVE = "evaluation_metrics/results.csv"
    LLM_MD_NAME = "extras/llm/LLM_Report.md"
    LLM_DOCX_NAME = "extras/llm/LLM_Report.docx"

    @to_obj(PostTrainContext)
    def __call__(self, ctx: PostTrainContext):
        project_path = ctx.project_path

        results_file = os.path.join(project_path, self.RESULTS_RELATIVE)
        llm_md_path = os.path.join(project_path, self.LLM_MD_NAME)
        llm_docx_path = os.path.join(project_path, self.LLM_DOCX_NAME)
        
        os.makedirs(os.path.dirname(llm_md_path), exist_ok=True)

        try:
            report = TrainingReportAnalyzer().analyze(results_file)
            
            # Save MD
            with open(llm_md_path, "w", encoding="utf-8") as f:
                f.write(report)
                
            # Save DOCX
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                from pathlib import Path
                
                doc = Document()
                media_dir = Path("/app/media")
                wtrain_img = media_dir / "wtrain.jpg"
                wpipe_img = media_dir / "wpipe.jpg"
                logo_img = media_dir / "logo.jpg"

                if wtrain_img.exists():
                    doc.add_picture(str(wtrain_img), width=Inches(6.0))
                    doc.add_paragraph("WTrain: Sistema completo de entrenamiento distribuido para modelos de Inteligencia Artificial.")
                    doc.add_page_break()

                if wpipe_img.exists():
                    doc.add_picture(str(wpipe_img), width=Inches(6.0))
                    doc.add_paragraph("El sistema hace uso de WPipe: un framework de pipelines profesional, rápido, eficiente y con características forenses avanzadas.")
                    doc.add_page_break()

                for line in report.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=0)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=2)
                    elif line.strip():
                        doc.add_paragraph(line)

                if logo_img.exists():
                    doc.add_page_break()
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(str(logo_img), width=Inches(3.0))
                    doc.add_paragraph("WTrain hace parte del paquete de la Wisrovi Suit, desarrollada por William Rodriguez (Wisrovi).").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                doc.save(llm_docx_path)
            except Exception as e:
                print(f"[LLMAnalyzer] Failed to generate DOCX: {e}")
                
            print(
                f"[LLMAnalyzer] Report written to {llm_md_path} and {llm_docx_path} "
                f"({len(report)} chars)."
            )
            return {"llm_report": report, "llm_md_path": llm_md_path, "llm_docx_path": llm_docx_path}
        except Exception as exc:
            print(f"[LLMAnalyzer] Failed: {exc}")
            return {"llm_report": "", "llm_md_path": "", "error": str(exc)}
