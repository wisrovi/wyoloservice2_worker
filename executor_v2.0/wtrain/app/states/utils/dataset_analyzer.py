import os
import cv2
import json
import yaml
import shutil
import hashlib
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
from pathlib import Path
from typing import Any
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Heuristics & Thresholds
CLASS_IMBALANCE_RATIO = 1.5
IMAGE_SIZE_VARIATION_THRESHOLD = 400
BBOX_AREA_VARIATION_RATIO = 20
ASPECT_RATIO_DIVERSITY_THRESHOLD = 50
BBOX_ASPECT_RATIO_VARIANCE = 2.0
CENTER_POSITION_BIN_COUNT = 75
BBOX_WH_CORRELATION_THRESHOLD = 0.3

class DatasetAnalyzer:
    """
    Analyze a dataset and automatically determine its type.
    This class detects whether the dataset is intended for
    classification, object detection, or segmentation tasks
    and delegates the analysis to the corresponding analyzer.
    """
    def analyze(self, dataset_path: str | Path) -> dict[str, Any]:
        dataset_path = Path(dataset_path)
        if dataset_path.is_file() and dataset_path.suffix in ['.yaml', '.yml']:
            dataset_path = dataset_path.parent
                
        dataset_type = self.detect_dataset_type(dataset_path)

        if dataset_type == "classification":
            stats = ClassificationAnalyzer().analyze(dataset_path)
        elif dataset_type == "detection":
            stats = DetectionAnalyzer().analyze(dataset_path)
        elif dataset_type == "segmentation":
            stats = SegmentationAnalyzer().analyze(dataset_path)
        else:
            return {"error": "Invalid dataset type"}
            
        try:
            EDAReportGenerator().generate_report(stats, dataset_path.name, dataset_type)
        except Exception as e:
            import traceback
            error_msg = f"Failed to generate EDA report: {e}\n{traceback.format_exc()}"
            print(error_msg)
            stats["eda_report_error"] = error_msg

        return stats

    def detect_dataset_type(self, dataset_path: str | Path) -> str:
        dataset_path = Path(dataset_path)
        train_path = dataset_path / "train"

        # Detection and segmentation (YOLO)
        if (train_path / "images").exists() and (train_path / "labels").exists():
            return self.detect_yolo_type(dataset_path)

        # Classification
        class_dirs = [item for item in dataset_path.iterdir() if item.is_dir()]
        if len(class_dirs) > 0:
            return "classification"

        return "unknown"

    def detect_yolo_type(self, dataset_path: str | Path) -> str:
        for split in ["train", "val", "test"]:
            labels_path = dataset_path / split / "labels"
            if not labels_path.exists():
                continue

            for label_file in labels_path.rglob("*.txt"):
                with open(label_file, encoding="utf-8") as f:
                    for line in f:
                        values = line.split()
                        if not values:
                            continue
                        if len(values) == 5:
                            return "detection"
                        if len(values) > 5:
                            return "segmentation"
        return "unknown"

    def load_class_names(self, dataset_path: str | Path) -> dict[str, str]:
        dataset_path = Path(dataset_path)
        yaml_files = list(dataset_path.glob("*.yaml"))
        
        if not yaml_files:
           yaml_path = dataset_path / "data.yaml"
           if os.path.exists(yaml_path):
                with open(yaml_path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f)
                    return data.get("names", {})
           return {}

        with open(yaml_files[0], encoding="utf-8") as f:
            data = yaml.safe_load(f)

        names = data.get("names", {})
        if isinstance(names, list):
            return {str(idx): name for idx, name in enumerate(names)}
        if isinstance(names, dict):
            return {str(idx): name for idx, name in names.items()}
        return {}


class DataQualityAnalyzer:
    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}

    def get_all_images(self, folder, split_name):
        return [
            (os.path.join(root, f), split_name)
            for root, _, files in os.walk(folder)
            for f in files if f.lower().endswith(tuple(self.IMAGE_EXTENSIONS))
        ]

    def detect_duplicates_and_overlaps(self, dataset_path: Path, dataset_type: str):
        image_hashes = {}
        duplicates = []
        image_paths = []

        if dataset_type == "classification":
            for split_folder in ["train", "test", "val"]:
                split_path = dataset_path / split_folder
                if split_path.is_dir():
                    image_paths.extend(self.get_all_images(str(split_path), split_folder))
        else:
            for split_folder in ["train", "test", "val"]:
                split_path = dataset_path / split_folder / "images"
                if split_path.is_dir():
                    image_paths.extend(self.get_all_images(str(split_path), split_folder))

        for img_path, split in image_paths:
            img = cv2.imread(img_path)
            if img is not None:
                img_hash = hashlib.sha1(img).hexdigest()
                if img_hash in image_hashes:
                    duplicates.append({
                        "img1": os.path.basename(img_path), "split1": split,
                        "img2": os.path.basename(image_hashes[img_hash][0]), "split2": image_hashes[img_hash][1]
                    })
                else:
                    image_hashes[img_hash] = (img_path, split)

        return duplicates

    def validate_image_quality(self, dataset_path: Path, dataset_type: str):
        corrupt_images = []
        small_images = []
        image_paths = []

        if dataset_type == "classification":
            for split_folder in ["train", "test", "val"]:
                split_path = dataset_path / split_folder
                if split_path.is_dir():
                    image_paths.extend([p[0] for p in self.get_all_images(str(split_path), split_folder)])
        else:
            for split_folder in ["train", "test", "val"]:
                split_path = dataset_path / split_folder / "images"
                if split_path.is_dir():
                    image_paths.extend([p[0] for p in self.get_all_images(str(split_path), split_folder)])

        for image_path in image_paths:
            img_name = os.path.basename(image_path)
            try:
                img = cv2.imread(image_path)
                if img is None:
                    corrupt_images.append(img_name)
                else:
                    height, width, _ = img.shape
                    if width < 64 or height < 64:
                        small_images.append({"name": img_name, "w": width, "h": height})
            except Exception:
                corrupt_images.append(img_name)

        return corrupt_images, small_images


class ClassificationAnalyzer:
    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}

    def analyze(self, dataset_path: str | Path) -> dict[str, Any]:
        dataset_path = Path(dataset_path)
        total_images = 0
        class_distribution = {}
        split_distribution = {}
        split_percentage_distribution = {}
        widths, heights = [], []
        aspect_ratios = []
        
        splits = ["train", "val", "test", "inference"]

        for split in splits:
            split_path = dataset_path / split
            if not split_path.exists():
                continue
            split_distribution[split] = {}

            for class_dir in split_path.iterdir():
                if not class_dir.is_dir():
                    continue
                image_count = 0

                for file in class_dir.rglob("*"):
                    if file.is_file() and file.suffix.lower() in self.IMAGE_EXTENSIONS:
                        image_count += 1
                        try:
                            with Image.open(file) as img:
                                widths.append(img.width)
                                heights.append(img.height)
                                if img.height > 0:
                                    aspect_ratios.append(img.width / img.height)
                        except:
                            pass

                split_distribution[split][class_dir.name] = image_count
                class_distribution[class_dir.name] = class_distribution.get(class_dir.name, 0) + image_count
                total_images += image_count

        for split, data in split_distribution.items():
            split_images = sum(data.values())
            if total_images > 0:
                split_percentage_distribution[split] = round(split_images * 100 / total_images, 2)

        dq = DataQualityAnalyzer()
        duplicates = dq.detect_duplicates_and_overlaps(dataset_path, "classification")
        corrupt, small = dq.validate_image_quality(dataset_path, "classification")

        return {
            "dataset_type": "classification",
            "num_classes": len(class_distribution),
            "total_images": total_images,
            "class_distribution": class_distribution,
            "split_distribution": split_distribution,
            "split_percentage_distribution": split_percentage_distribution,
            "image_dimensions": {"widths": widths, "heights": heights},
            "aspect_ratios": aspect_ratios,
            "data_quality": {
                "duplicates": duplicates,
                "corrupt": corrupt,
                "small": small
            },
            "dataset_path": str(dataset_path)
        }


class DetectionAnalyzer:
    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}

    def analyze(self, dataset_path: str | Path) -> dict[str, Any]:
        dataset_path = Path(dataset_path)
        total_images = 0
        total_annotations = 0
        class_names = DatasetAnalyzer().load_class_names(dataset_path)
        
        class_distribution = {}
        split_distribution = {}
        split_percentage_distribution = {}
        
        widths, heights = [], []
        aspect_ratios = []
        bbox_areas = []
        bbox_aspect_ratios = []
        x_centers, y_centers = [], []
        bbox_widths, bbox_heights = [], []

        splits = ["train", "val", "test", "inference"]

        for split in splits:
            split_path = dataset_path / split
            if not split_path.exists():
                continue

            images_split_path = dataset_path / split / "images"
            labels_split_path = dataset_path / split / "labels"

            split_distribution[split] = {"images": 0, "annotations": 0, "classes": {}}

            for image_file in images_split_path.rglob("*"):
                if image_file.suffix.lower() in self.IMAGE_EXTENSIONS:
                    total_images += 1
                    split_distribution[split]["images"] += 1
                    try:
                        with Image.open(image_file) as img:
                            widths.append(img.width)
                            heights.append(img.height)
                            if img.height > 0:
                                aspect_ratios.append(img.width / img.height)
                    except:
                        pass

            for label_file in labels_split_path.rglob("*.txt"):
                with open(label_file, "r", encoding="utf-8") as f:
                    for line in f:
                        values = line.split()
                        if not values:
                            continue
                        
                        total_annotations += 1
                        split_distribution[split]["annotations"] += 1
                        
                        if len(values) >= 5:
                            try:
                                _, x_c, y_c, b_w, b_h = map(float, values[:5])
                                bbox_areas.append(b_w * b_h)
                                x_centers.append(x_c)
                                y_centers.append(y_c)
                                bbox_widths.append(b_w)
                                bbox_heights.append(b_h)
                                if b_h > 0:
                                    bbox_aspect_ratios.append(b_w / b_h)
                            except ValueError:
                                pass

                        class_id = values[0]
                        class_name = class_names.get(class_id, class_id)
                        split_distribution[split]["classes"][class_name] = split_distribution[split]["classes"].get(class_name, 0) + 1
                        class_distribution[class_name] = class_distribution.get(class_name, 0) + 1

        for split, data in split_distribution.items():
            if total_images > 0:
                split_percentage_distribution[split] = round(data["images"] * 100 / total_images, 2)

        dq = DataQualityAnalyzer()
        duplicates = dq.detect_duplicates_and_overlaps(dataset_path, "detection")
        corrupt, small = dq.validate_image_quality(dataset_path, "detection")

        return {
            "dataset_type": "detection",
            "total_images": total_images,
            "total_annotations": total_annotations,
            "num_classes": len(class_distribution),
            "class_distribution": class_distribution,
            "split_distribution": split_distribution,
            "split_percentage_distribution": split_percentage_distribution,
            "image_dimensions": {"widths": widths, "heights": heights},
            "aspect_ratios": aspect_ratios,
            "bbox_areas": bbox_areas,
            "bbox_aspect_ratios": bbox_aspect_ratios,
            "bbox_centers": {"x": x_centers, "y": y_centers},
            "bbox_dims": {"w": bbox_widths, "h": bbox_heights},
            "data_quality": {
                "duplicates": duplicates,
                "corrupt": corrupt,
                "small": small
            },
            "dataset_path": str(dataset_path)
        }


class SegmentationAnalyzer(DetectionAnalyzer):
    def analyze(self, dataset_path: str | Path) -> dict[str, Any]:
        res = super().analyze(dataset_path)
        res["dataset_type"] = "segmentation"
        return res


class EDAReportGenerator:
    def __init__(self, results_dir="/wyolo/worker/train_service_results/extras/eda"):
        self.results_dir = Path(results_dir)
        # Limpieza profunda de los artefactos EDA antiguos antes de generar nuevos
        if self.results_dir.exists():
            shutil.rmtree(self.results_dir)
        self.results_dir.mkdir(parents=True, exist_ok=True)
        self.plots_dir = self.results_dir / "plots"
        self.plots_dir.mkdir(exist_ok=True)

    def add_index(self, doc, section_titles):
        index_doc = Document()
        index_doc.add_page_break()
        index_doc.add_paragraph("Índice", style="Heading 1")
        for i, title in enumerate(section_titles, 1):
            index_doc.add_paragraph(f"{i}. {title}", style="Normal")
        index_doc.add_page_break()

        index_elements = [p._element for p in index_doc.paragraphs]
        if index_doc.paragraphs:
            index_elements.append(index_doc.paragraphs[-1]._element.getparent())

        insert_pos = 8
        body = doc._body._element
        for element in reversed(index_elements):
            body.insert(insert_pos, element)

    def generate_report(self, stats: dict, dataset_name: str, dataset_type: str):
        md_content = f"# EDA Report: {dataset_name}\n\n"
        doc = Document()
        section_titles = []

        # Portada
        doc.add_paragraph(f"Análisis Exploratorio de Datos (EDA)", style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("Informe de Análisis Exploratorio de Datos", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("\n", style="Normal")
        doc.add_paragraph(f"Proyecto/Dataset: {dataset_name}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Tipo de Dataset: {dataset_type.capitalize()}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph("\n", style="Normal")

        # 1. Calidad de Datos
        section_titles.append("Validación de Calidad de Imágenes")
        doc.add_paragraph("Validación de Calidad de Imágenes:", style="Heading 1")
        dq = stats.get("data_quality", {})
        corrupt = dq.get("corrupt", [])
        small = dq.get("small", [])
        dups = dq.get("duplicates", [])
        
        md_content += "## Validación de Calidad de Imágenes\n\n"
        
        if corrupt:
            doc.add_paragraph("Imágenes corruptas encontradas:", style="Heading 2")
            md_content += "### Imágenes corruptas\n"
            for img in corrupt[:20]:
                doc.add_paragraph(f"- {img}")
                md_content += f"- {img}\n"
            if len(corrupt) > 20:
                doc.add_paragraph(f"... y {len(corrupt)-20} más.")
                md_content += f"... y {len(corrupt)-20} más.\n"
        else:
            doc.add_paragraph("No se encontraron imágenes corruptas.")
            md_content += "No se encontraron imágenes corruptas.\n\n"

        if small:
            doc.add_paragraph("Imágenes con dimensiones pequeñas encontradas (<64px):", style="Heading 2")
            md_content += "### Imágenes diminutas\n"
            for img in small[:20]:
                doc.add_paragraph(f"- {img['name']} ({img['w']}x{img['h']})")
                md_content += f"- {img['name']} ({img['w']}x{img['h']})\n"
            if len(small) > 20:
                doc.add_paragraph(f"... y {len(small)-20} más.")
                md_content += f"... y {len(small)-20} más.\n"
        else:
            doc.add_paragraph("No se encontraron imágenes con dimensiones pequeñas.")
            md_content += "No se encontraron imágenes con dimensiones pequeñas.\n\n"

        if dups:
            doc.add_paragraph("Imágenes Duplicadas:", style="Heading 2")
            md_content += "### Imágenes duplicadas\n"
            for dup in dups[:20]:
                msg = f"- '{dup['img1']}' en '{dup['split1']}' y '{dup['img2']}' en '{dup['split2']}'"
                doc.add_paragraph(msg)
                md_content += f"{msg}\n"
            if len(dups) > 20:
                doc.add_paragraph(f"... y {len(dups)-20} más.")
                md_content += f"... y {len(dups)-20} más.\n"
        else:
            doc.add_paragraph("No se encontraron imágenes duplicadas.")
            md_content += "No se encontraron imágenes duplicadas.\n\n"

        doc.add_page_break()

        # 2. Distribución de Clases
        if "class_distribution" in stats:
            classes = list(stats["class_distribution"].keys())
            counts = list(stats["class_distribution"].values())
            if classes:
                plt.figure(figsize=(12, max(6, len(classes)*0.3)))
                sns.barplot(x=counts, y=classes, palette="viridis")
                plt.title("Distribución de Clases")
                plt.xlabel("Instancias")
                plt.ylabel("Clase")
                plt.tight_layout()
                dist_path = self.plots_dir / "class_distribution.png"
                plt.savefig(dist_path)
                plt.close()

                md_content += "## Distribución de Clases\n\n![Class Distribution](plots/class_distribution.png)\n\n"
                doc.add_heading("Distribución de Clases", level=1)
                section_titles.append("Distribución de Clases")
                doc.add_picture(str(dist_path), width=Inches(6.0))

        # 3. Splits
        if "split_percentage_distribution" in stats:
            splits = list(stats["split_percentage_distribution"].keys())
            percentages = list(stats["split_percentage_distribution"].values())
            if sum(percentages) > 0:
                plt.figure(figsize=(8, 8))
                plt.pie(percentages, labels=splits, autopct="%1.1f%%", startangle=140, colors=sns.color_palette("pastel"))
                plt.title("Distribución de Splits")
                split_path = self.plots_dir / "split_distribution.png"
                plt.savefig(split_path)
                plt.close()

                md_content += "## Distribución de Splits\n\n![Split Distribution](plots/split_distribution.png)\n\n"
                doc.add_heading("Distribución de Splits", level=1)
                section_titles.append("Distribución de Splits")
                doc.add_picture(str(split_path), width=Inches(5.0))

        doc.add_page_break()

        # 4. Dimensiones de Imágenes
        dims = stats.get("image_dimensions", {})
        if dims.get("widths") and dims.get("heights"):
            plt.figure(figsize=(10, 6))
            sns.histplot(dims["widths"], bins=30, color="blue", label="Width", alpha=0.6)
            sns.histplot(dims["heights"], bins=30, color="red", label="Height", alpha=0.6)
            plt.xlabel("Píxeles")
            plt.ylabel("Frecuencia")
            plt.title("Distribución de Tamaños de Imágenes")
            plt.legend()
            sz_path = self.plots_dir / "image_size_distribution.png"
            plt.savefig(sz_path)
            plt.close()
            
            md_content += "## Tamaños de Imágenes\n\n![Image Sizes](plots/image_size_distribution.png)\n\n"
            doc.add_heading("Distribución de Tamaños de Imágenes", level=1)
            section_titles.append("Distribución de Tamaños de Imágenes")
            doc.add_picture(str(sz_path), width=Inches(6.0))

        # 5. Aspect Ratio (Imágenes)
        aspect = stats.get("aspect_ratios", [])
        if aspect:
            plt.figure(figsize=(10, 6))
            sns.histplot(aspect, bins=30, color="purple")
            plt.xlabel("Aspect Ratio (Width/Height)")
            plt.ylabel("Frecuencia")
            plt.title("Relación de Aspecto de Imágenes")
            ar_path = self.plots_dir / "aspect_ratio_distribution.png"
            plt.savefig(ar_path)
            plt.close()

            md_content += "## Relación de Aspecto de Imágenes\n\n![Aspect Ratio](plots/aspect_ratio_distribution.png)\n\n"
            doc.add_heading("Distribución de Relación de Aspecto", level=1)
            section_titles.append("Distribución de Relación de Aspecto")
            doc.add_picture(str(ar_path), width=Inches(6.0))

        doc.add_page_break()

        # 6. Geometría de Bounding Boxes (Detección/Segmentación)
        if stats.get("dataset_type") in ["detection", "segmentation"]:
            # Áreas
            if stats.get("bbox_areas"):
                plt.figure(figsize=(10, 6))
                sns.histplot(stats["bbox_areas"], bins=50, kde=True, color="green")
                plt.title("Distribución de Áreas de Bounding Boxes (Normalizadas)")
                plt.xlabel("Área (width * height)")
                plt.ylabel("Frecuencia")
                plt.tight_layout()
                bbox_path = self.plots_dir / "bbox_areas.png"
                plt.savefig(bbox_path)
                plt.close()

                md_content += "## Bounding Box Áreas\n\n![BBox Areas](plots/bbox_areas.png)\n\n"
                doc.add_heading("Distribución de Áreas de BBox", level=1)
                section_titles.append("Distribución de Áreas de BBox")
                doc.add_picture(str(bbox_path), width=Inches(6.0))

            # Aspect Ratio BBOX
            if stats.get("bbox_aspect_ratios"):
                plt.figure(figsize=(10, 6))
                sns.histplot(stats["bbox_aspect_ratios"], bins=30, color="orange")
                plt.title("Distribución de Relación de Aspecto de BBox")
                plt.xlabel("Aspect Ratio (W/H)")
                bbox_ar_path = self.plots_dir / "bbox_aspect_ratios.png"
                plt.savefig(bbox_ar_path)
                plt.close()
                doc.add_heading("Aspect Ratio de Bounding Boxes", level=1)
                section_titles.append("Aspect Ratio de BBox")
                doc.add_picture(str(bbox_ar_path), width=Inches(6.0))

            # Centros
            centers = stats.get("bbox_centers", {})
            if centers.get("x") and centers.get("y"):
                plt.figure(figsize=(12, 6))
                plt.subplot(1, 2, 1)
                sns.histplot(centers["x"], bins=30, color="purple")
                plt.title("Centros X")
                plt.subplot(1, 2, 2)
                sns.histplot(centers["y"], bins=30, color="blue")
                plt.title("Centros Y")
                cen_path = self.plots_dir / "bbox_centers.png"
                plt.savefig(cen_path)
                plt.close()
                doc.add_heading("Posiciones Centrales de BBox", level=1)
                section_titles.append("Posiciones Centrales de BBox")
                doc.add_picture(str(cen_path), width=Inches(6.0))

            # Scatter W vs H
            dims = stats.get("bbox_dims", {})
            if dims.get("w") and dims.get("h"):
                plt.figure(figsize=(8, 8))
                plt.scatter(dims["w"], dims["h"], alpha=0.3)
                plt.title("BBox Width vs Height")
                plt.xlabel("Width")
                plt.ylabel("Height")
                sc_path = self.plots_dir / "bbox_scatter.png"
                plt.savefig(sc_path)
                plt.close()
                doc.add_heading("Dispersión Ancho vs Alto de BBox", level=1)
                section_titles.append("Dispersión Ancho vs Alto")
                doc.add_picture(str(sc_path), width=Inches(6.0))

        doc.add_page_break()

        # 7. Conclusiones Heurísticas
        doc.add_heading("Conclusiones del Análisis", level=1)
        section_titles.append("Conclusiones del Análisis")
        md_content += "## Conclusiones del Análisis\n\n"
        
        conclusions = []
        
        # Corruptas
        if corrupt:
            conclusions.append("Se detectaron imágenes corruptas que deben ser eliminadas o reemplazadas.")
        else:
            conclusions.append("No se encontraron imágenes corruptas.")
            
        # Pequeñas
        if small:
            conclusions.append("Existen imágenes con dimensiones muy pequeñas que podrían afectar el rendimiento.")
        else:
            conclusions.append("Las dimensiones mínimas de las imágenes son adecuadas.")

        # Duplicados
        if dups:
            conclusions.append("Se encontraron imágenes duplicadas o superposiciones entre splits, cuidado con el data leakage.")
        else:
            conclusions.append("No se detectaron duplicados problemáticos.")

        # Clases
        if "class_distribution" in stats and stats["class_distribution"]:
            counts = list(stats["class_distribution"].values())
            max_c, min_c = max(counts), min(counts)
            if (max_c / (min_c + 1e-5)) > CLASS_IMBALANCE_RATIO:
                conclusions.append("La distribución de clases está fuertemente desbalanceada. Considere recolectar más datos para clases minoritarias o aplicar pesos.")
            else:
                conclusions.append("La distribución de clases es equilibrada.")
                
        # Dimensions
        if dims.get("widths"):
            w, h = dims["widths"], dims["heights"]
            if max(w) - min(w) > IMAGE_SIZE_VARIATION_THRESHOLD or max(h) - min(h) > IMAGE_SIZE_VARIATION_THRESHOLD:
                conclusions.append("Alta variabilidad en los tamaños de las imágenes. El modelo requerirá padding o resize agresivo.")
            else:
                conclusions.append("Los tamaños de las imágenes son uniformes.")

        # Bbox Areas
        if stats.get("bbox_areas"):
            areas = stats["bbox_areas"]
            if max(areas) / (min(areas) + 1e-5) > BBOX_AREA_VARIATION_RATIO:
                conclusions.append("Mucha variabilidad en las áreas de BBox. El modelo detectará objetos de múltiples tamaños.")
            else:
                conclusions.append("Las áreas de objetos son relativamente uniformes.")

        for line in conclusions:
            doc.add_paragraph(line, style="List Number")
            md_content += f"1. {line}\n"

        try:
            self.add_index(doc, section_titles)
        except:
            pass

        # Save MD
        md_file = self.results_dir / "EDA_Report.md"
        with open(md_file, "w") as f:
            f.write(md_content)

        # Save DOCX
        docx_file = self.results_dir / "EDA_Report.docx"
        doc.save(str(docx_file))

        return str(self.results_dir)


class DatasetEDAState:
    def execute(self, dataset_path: str | Path) -> dict[str, Any]:
        dataset_path = Path(dataset_path)
        if not dataset_path.exists():
            raise FileNotFoundError(f"Dataset not found: {dataset_path}")
        return DatasetAnalyzer().analyze(dataset_path)

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        raise ValueError("Usage: python dataset_analyzer.py <dataset_path>")
    dataset_path = Path(sys.argv[1])
    if not dataset_path.exists():
        raise FileNotFoundError(f"Dataset not found: {dataset_path}")
    result = DatasetAnalyzer().analyze(dataset_path)
    print(json.dumps(result, indent=4))
